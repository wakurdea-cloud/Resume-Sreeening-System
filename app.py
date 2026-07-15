import streamlit as st
import joblib
import pdfplumber
from docx import Document
import numpy as np

# -----------------------
# Page Config
# -----------------------
st.set_page_config(
    page_title="AI Resume Screening System",
    page_icon="📄",
    layout="wide"
)

# -----------------------
# Load Model
# -----------------------
model = joblib.load("resume_screening_model.pkl")

# -----------------------
# CSS
# -----------------------
st.markdown("""
<style>
.main{
    background-color:#f5f7fa;
}
h1{
    color:#1f4e79;
    text-align:center;
}
.stButton>button{
    width:100%;
    background:#1f77b4;
    color:white;
    border-radius:10px;
    height:50px;
    font-size:18px;
}
</style>
""", unsafe_allow_html=True)

st.title("📄 AI Resume Screening System")

st.write("Upload your Resume and click **Predict Resume**.")

# -----------------------
# File Upload
# -----------------------

uploaded_file = st.file_uploader(
    "Upload Resume",
    type=["pdf", "docx", "txt"]
)

resume_text = ""

if uploaded_file is not None:

    # प्रत्येक नवीन upload साठी file pointer reset
    uploaded_file.seek(0)

    if uploaded_file.name.lower().endswith(".pdf"):

        with pdfplumber.open(uploaded_file) as pdf:
            st.write("Resume Text Length:", len(resume_text))
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

        resume_text = "\n".join(pages)

    elif uploaded_file.name.lower().endswith(".docx"):

        doc = Document(uploaded_file)
        resume_text = "\n".join([p.text for p in doc.paragraphs])

    elif uploaded_file.name.lower().endswith(".txt"):

        resume_text = uploaded_file.read().decode("utf-8", errors="ignore")

    st.success(f"Uploaded: {uploaded_file.name}")

resume_text = ""

if uploaded_file:

    if uploaded_file.name.endswith(".pdf"):

        with pdfplumber.open(uploaded_file) as pdf:

            for page in pdf.pages:

                text = page.extract_text()

                if text:
                    resume_text += text + "\n"

    elif uploaded_file.name.endswith(".docx"):

        doc = Document(uploaded_file)

        for para in doc.paragraphs:

            resume_text += para.text + "\n"

    elif uploaded_file.name.endswith(".txt"):

        resume_text = uploaded_file.read().decode("utf-8")

# -----------------------
# Preview
# -----------------------

if resume_text:

    st.subheader("📄 Resume Preview")

    st.text_area(
        "",
        value=resume_text,
        height=250
    )

# -----------------------
# Prediction
# -----------------------

if st.button("🚀 Predict Resume"):

    if not resume_text:

        st.warning("Please upload a Resume.")

    else:

        prediction = model.predict([resume_text])[0]

        probabilities = model.predict_proba([resume_text])[0]

        confidence = probabilities.max()*100

        st.success(f"✅ Predicted Category : **{prediction}**")

        st.subheader("🎯 Match Percentage")

        st.progress(int(confidence))

        st.write(f"### {confidence:.2f}%")

        st.subheader("🏆 Top 3 Matching Categories")

        classes = model.classes_

        top3 = np.argsort(probabilities)[::-1][:3]

        for i in top3:

            st.write(
                f"**{classes[i]}** : {probabilities[i]*100:.2f}%"
            )

        st.balloons()

# -----------------------
# Analyze Another Resume
# -----------------------

